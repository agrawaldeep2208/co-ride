"""
Convert README.md to a professional Word document with rendered Mermaid diagrams
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import json
import os
from io import BytesIO
import subprocess
import shutil

# Mermaid diagrams to render
DIAGRAMS = {
    "architecture": """graph TB
    subgraph Frontend["Frontend (React + TypeScript)"]
        Pages["14+ Pages<br/>- Landing<br/>- Auth<br/>- Ride Management<br/>- Notifications<br/>- Admin Dashboard"]
        Maps["Leaflet Maps<br/>Geolocation"]
        Auth["Auth Context<br/>JWT/OAuth"]
    end
    
    subgraph Backend["Backend (Express.js + Node.js)"]
        Routes["8 Route Modules<br/>- Auth<br/>- Ride<br/>- User<br/>- Vehicle<br/>- Payment<br/>- Rating<br/>- Notification<br/>- Admin"]
        Auth2["Authentication<br/>- JWT<br/>- OTP (Email)<br/>- Google OAuth"]
        Middleware["Middleware<br/>- Protected Routes<br/>- Role-based Access<br/>- File Upload"]
    end
    
    subgraph Database["MongoDB"]
        Collections["9 Collections<br/>- User<br/>- Ride<br/>- RideRequest<br/>- Vehicle<br/>- Payment<br/>- Rating<br/>- Notification<br/>- OTP<br/>- UserStats"]
    end
    
    subgraph External["External Services"]
        Gmail["Gmail SMTP<br/>Nodemailer"]
        Google["Google OAuth<br/>Authentication"]
        Multer["Multer<br/>File Upload"]
    end
    
    Frontend -->|API Calls| Backend
    Backend -->|CRUD Operations| Database
    Backend -->|Email/OTP| Gmail
    Frontend -->|OAuth| Google
    Backend -->|License Upload| Multer
    
    style Frontend fill:#3498db,stroke:#2c3e50,color:#fff
    style Backend fill:#e74c3c,stroke:#2c3e50,color:#fff
    style Database fill:#2ecc71,stroke:#2c3e50,color:#fff
    style External fill:#f39c12,stroke:#2c3e50,color:#fff""",
    
    "database": """graph LR
    USER["👤 USER<br/>- fullName<br/>- email<br/>- phone"]
    RIDE["🚗 RIDE<br/>- source<br/>- destination<br/>- pricePerSeat"]
    VEHICLE["🛞 VEHICLE<br/>- vehicleModel<br/>- vehicleNumber"]
    RIDEREQ["📋 RIDEREQUEST<br/>- seatsRequested<br/>- status"]
    PAYMENT["💳 PAYMENT<br/>- amount<br/>- status"]
    RATING["⭐ RATING<br/>- rating 1-5<br/>- feedback"]
    NOTIF["🔔 NOTIFICATION<br/>- type<br/>- message"]
    STATS["📊 USERSTATS<br/>- averageRating<br/>- totalRides"]
    OTP["🔐 OTP<br/>- code<br/>- expiry"]
    
    USER -->|creates| RIDE
    USER -->|owns| VEHICLE
    USER -->|makes| RIDEREQ
    USER -->|pays| PAYMENT
    USER -->|gives| RATING
    USER -->|receives| NOTIF
    USER -->|has| STATS
    USER -->|verifies| OTP
    
    RIDE -->|uses| VEHICLE
    RIDE -->|has| RIDEREQ
    RIDE -->|includes| PAYMENT
    RIDE -->|gets| RATING
    
    style USER fill:#3498db,stroke:#2c3e50,color:#fff
    style RIDE fill:#e74c3c,stroke:#2c3e50,color:#fff
    style VEHICLE fill:#2ecc71,stroke:#2c3e50,color:#fff
    style RIDEREQ fill:#f39c12,stroke:#2c3e50,color:#fff
    style PAYMENT fill:#9b59b6,stroke:#2c3e50,color:#fff
    style RATING fill:#1abc9c,stroke:#2c3e50,color:#fff
    style NOTIF fill:#e67e22,stroke:#2c3e50,color:#fff
    style STATS fill:#34495e,stroke:#2c3e50,color:#fff
    style OTP fill:#c0392b,stroke:#2c3e50,color:#fff"""
}

def render_mermaid_diagram(diagram_code, filename):
    """Render Mermaid diagram to PNG using quickchart.io service"""
    print(f"Rendering {filename}...")
    
    try:
        # Create diagrams folder
        os.makedirs("diagrams", exist_ok=True)
        
        # Method 1: Try quickchart.io (more reliable for rendering)
        import base64
        encoded = base64.b64encode(diagram_code.encode()).decode()
        url = f"https://quickchart.io/mermaid?mermaid={encoded}"
        
        response = requests.get(url, timeout=15)
        
        if response.status_code == 200:
            img_path = f"diagrams/{filename}.png"
            with open(img_path, 'wb') as f:
                f.write(response.content)
            print(f"✓ Successfully rendered {filename}")
            return img_path
        
        # Method 2: Try mermaid.ink as fallback
        encoded_ink = base64.b64encode(diagram_code.encode()).decode()
        url_ink = f"https://mermaid.ink/img/{encoded_ink}"
        response = requests.get(url_ink, timeout=15)
        
        if response.status_code == 200:
            img_path = f"diagrams/{filename}.png"
            with open(img_path, 'wb') as f:
                f.write(response.content)
            print(f"✓ Successfully rendered {filename} (mermaid.ink)")
            return img_path
        else:
            print(f"✗ Failed to render {filename}: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"✗ Error rendering {filename}: {e}")
        return None

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    doc.add_heading(text, level=level)

def add_bullet_point(doc, text, level=0):
    """Add bullet points to the document"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))

def set_table_header_style(row):
    """Set table header style"""
    for cell in row.cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

def create_word_document():
    """Create the Word document"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('coRide - Full-Stack Ride-Sharing Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('A Modern Carpooling Application Built with MERN Stack')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    doc.add_paragraph()  # Spacer
    
    # Project Overview
    add_heading(doc, '🎯 Project Overview', 1)
    doc.add_paragraph('coRide is a peer-to-peer ride-sharing platform that enables users to:')
    add_bullet_point(doc, 'Create and manage rides')
    add_bullet_point(doc, 'Search and join available rides')
    add_bullet_point(doc, 'Process payments securely')
    add_bullet_point(doc, 'Rate and review fellow users')
    add_bullet_point(doc, 'Track ride history and statistics')
    
    # Architecture Section
    add_heading(doc, '🏗️ Architecture', 1)
    
    add_heading(doc, 'Backend (Express.js + Node.js)', 2)
    add_bullet_point(doc, '32+ RESTful API endpoints across 8 modules')
    add_bullet_point(doc, 'JWT + OTP + Google OAuth authentication')
    add_bullet_point(doc, 'Role-based access control (RBAC)')
    add_bullet_point(doc, 'MongoDB database with 9 interconnected collections')
    
    add_heading(doc, 'Frontend (React + TypeScript)', 2)
    add_bullet_point(doc, '14+ responsive pages with Vite build optimization')
    add_bullet_point(doc, 'Leaflet maps for geolocation')
    add_bullet_point(doc, 'Real-time admin dashboard')
    add_bullet_point(doc, 'Tailwind CSS responsive design')
    
    # System Architecture Diagram
    add_heading(doc, 'System Architecture Diagram', 2)
    arch_img = render_mermaid_diagram(DIAGRAMS["architecture"], "architecture_diagram")
    if arch_img and os.path.exists(arch_img):
        try:
            doc.add_picture(arch_img, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Inserted {arch_img} into document")
        except Exception as e:
            print(f"✗ Could not insert image: {e}")
            doc.add_paragraph("(See README.md for System Architecture diagram)")
    else:
        doc.add_paragraph("(See README.md for System Architecture diagram)")
    
    doc.add_paragraph()  # Spacer
    
    # Database Schema
    add_heading(doc, '🗄️ Database Schema', 1)
    doc.add_paragraph('coRide utilizes 9 interconnected MongoDB collections:')
    
    # Database Collections Table
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Collection'
    header_cells[1].text = 'Purpose'
    set_table_header_style(table.rows[0])
    
    collections_data = [
        ('User', 'User authentication & profiles'),
        ('Ride', 'Ride creation & management'),
        ('RideRequest', 'Booking requests (Pending→Approved→Joined)'),
        ('Vehicle', 'User vehicle registration'),
        ('Payment', 'Transaction tracking'),
        ('Rating', 'User-to-user ratings (1-5 stars)'),
        ('Notification', 'Real-time system alerts'),
        ('UserStats', 'User performance metrics'),
        ('OTP', 'Email verification (5-min expiry)'),
    ]
    
    for i, (collection, purpose) in enumerate(collections_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = collection
        row_cells[1].text = purpose
    
    doc.add_paragraph()  # Spacer
    
    # Entity Relationship Diagram
    add_heading(doc, 'Entity Relationship Diagram', 2)
    db_img = render_mermaid_diagram(DIAGRAMS["database"], "database_diagram")
    if db_img and os.path.exists(db_img):
        try:
            doc.add_picture(db_img, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Inserted {db_img} into document")
        except Exception as e:
            print(f"✗ Could not insert image: {e}")
            doc.add_paragraph("(See README.md for Entity Relationship diagram)")
    else:
        doc.add_paragraph("(See README.md for Entity Relationship diagram)")
    
    doc.add_page_break()
    
    # Authentication & Security
    add_heading(doc, '🔐 Authentication & Security', 1)
    add_bullet_point(doc, 'JWT Tokens - 30-day expiration')
    add_bullet_point(doc, 'OTP Email Verification - 5-minute validity')
    add_bullet_point(doc, 'Google OAuth 2.0 - One-click login')
    add_bullet_point(doc, 'Password Hashing - bcryptjs (10 salt rounds)')
    add_bullet_point(doc, 'Role-Based Access Control - User & Admin roles')
    add_bullet_point(doc, 'File Validation - License upload with JPEG/PNG/JPG restriction')
    add_bullet_point(doc, 'Indian Vehicle Plate Validation - Regex pattern matching')
    
    # API Endpoints
    add_heading(doc, '📡 API Endpoints (32+)', 1)
    
    endpoint_sections = {
        'Authentication (5 endpoints)': [
            'POST /api/auth/send-otp - Send OTP to email',
            'POST /api/auth/register - Register with OTP verification',
            'POST /api/auth/login - Login credentials',
            'POST /api/auth/google-login - Google OAuth',
        ],
        'Rides (10+ endpoints)': [
            'POST /api/ride - Create ride',
            'GET /api/ride - Get all active rides',
            'GET /api/ride/created - User\'s created rides',
            'POST /api/ride/:id/request - Request to join',
            'PUT /api/ride/request/:id - Approve/reject request',
            'PUT /api/ride/:id/start - Start ride',
        ],
        'Vehicles (3 endpoints)': [
            'POST /api/vehicle - Register vehicle',
            'GET /api/vehicle/my-vehicles - User\'s vehicles',
            'DELETE /api/vehicle/:id - Delete vehicle',
        ],
        'Payments (2 endpoints)': [
            'GET /api/payment/ride/:rideId - Get ride payments',
            'PUT /api/payment/:id/pay - Complete payment',
        ],
        'Ratings (2 endpoints)': [
            'POST /api/rating - Submit rating (prevents duplicates)',
            'GET /api/rating/ride/:rideId - Get ride ratings',
        ],
        'Notifications (2+ endpoints)': [
            'GET /api/notification - Get all notifications',
            'PUT /api/notification/:id/read - Mark as read',
            'PUT /api/notification/read-all - Mark all as read',
        ],
        'Admin (3 endpoints)': [
            'GET /api/admin/stats - Dashboard statistics',
            'GET /api/admin/users - All users with stats',
            'GET /api/admin/rides - All rides',
        ],
    }
    
    for section_title, endpoints in endpoint_sections.items():
        add_heading(doc, section_title, 2)
        for endpoint in endpoints:
            add_bullet_point(doc, endpoint)
    
    doc.add_page_break()
    
    # Tech Stack
    add_heading(doc, '🛠️ Tech Stack', 1)
    
    add_heading(doc, 'Backend', 2)
    add_bullet_point(doc, 'Express.js 5.2.1')
    add_bullet_point(doc, 'Node.js')
    add_bullet_point(doc, 'MongoDB + Mongoose 9.3.0')
    add_bullet_point(doc, 'JWT Authentication')
    add_bullet_point(doc, 'Nodemailer (Email)')
    add_bullet_point(doc, 'Multer (File Upload)')
    add_bullet_point(doc, 'bcryptjs (Password Hashing)')
    
    add_heading(doc, 'Frontend', 2)
    add_bullet_point(doc, 'React 18.3.1')
    add_bullet_point(doc, 'TypeScript 5.5.3')
    add_bullet_point(doc, 'Vite 5.4.2')
    add_bullet_point(doc, 'Tailwind CSS 3.4.1')
    add_bullet_point(doc, 'React Router DOM 7.13.0')
    add_bullet_point(doc, 'Leaflet 1.9.4 (Maps)')
    add_bullet_point(doc, 'lucide-react (Icons)')
    
    # Key Features
    add_heading(doc, '🚀 Key Features', 1)
    add_bullet_point(doc, 'Real-Time Ride Management - Create, search, and join rides')
    add_bullet_point(doc, 'Multi-Method Authentication - JWT, OTP, Google OAuth')
    add_bullet_point(doc, 'Secure Payments - Wallet-based transaction system')
    add_bullet_point(doc, 'Rating System - 1-5 star ratings with feedback')
    add_bullet_point(doc, 'File Upload - License validation and storage')
    add_bullet_point(doc, 'Email Notifications - Ride updates and alerts')
    add_bullet_point(doc, 'Admin Dashboard - Real-time analytics and management')
    add_bullet_point(doc, 'Geolocation Maps - Leaflet integration for ride tracking')
    add_bullet_point(doc, 'Responsive Design - Mobile-friendly UI with Tailwind CSS')
    add_bullet_point(doc, 'Role-Based Access - User and Admin permissions')
    
    doc.add_page_break()
    
    # Resume Points
    add_heading(doc, '🎯 Resume Highlights', 1)
    
    resume_table = doc.add_table(rows=4, cols=1)
    resume_table.style = 'Light Grid Accent 1'
    
    resume_points = [
        'Developed 32+ REST APIs across 8 modules with complex multi-role request workflows and ride lifecycle management.',
        'Architected 9 MongoDB collections with JWT + OTP + Google OAuth authentication and role-based access control (RBAC).',
        'Built 14+ React-TypeScript pages with real-time admin dashboard, Leaflet maps, Multer file uploads, and Nodemailer integration.',
    ]
    
    for i, point in enumerate(resume_points):
        resume_table.rows[i+1].cells[0].text = point
    
    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Built with ❤️ - A Full-Stack MERN Project\ncoRide © 2026')
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    
    return doc

def main():
    print("=" * 60)
    print("📄 Converting README to Word Document with Diagrams")
    print("=" * 60)
    
    # Render diagrams
    print("\n🎨 Rendering Mermaid Diagrams...")
    for diagram_name, diagram_code in DIAGRAMS.items():
        render_mermaid_diagram(diagram_code, diagram_name + "_diagram")
    
    # Create document
    print("\n📝 Creating Word document...")
    doc = create_word_document()
    
    # Save document
    output_path = "coRide_Project_Report.docx"
    doc.save(output_path)
    
    print(f"\n✓ Document created successfully!")
    print(f"📍 Saved to: {os.path.abspath(output_path)}")
    print("=" * 60)

if __name__ == "__main__":
    main()
